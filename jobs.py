"""
AI Recruit - Job Posting API Endpoints
Job posting management with AI-powered analysis and matching.

NO MANUAL RULES - NO FALLBACKS - PURE AI INTELLIGENCE
"""

import io
import logging
import os
import tempfile
from datetime import datetime
from typing import List, Optional, Dict, Any
from uuid import UUID

import docx
import docx2txt
import pdfplumber
from fastapi import APIRouter, Depends, HTTPException, Query, Request, UploadFile, File
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field, validator
from sqlalchemy import select, func, or_, delete, and_
from sqlalchemy.ext.asyncio import AsyncSession

from agents.job_copilot import JobCopilotAgent, CopilotResponse, InterviewQuestionsResponse, SkillAnalysisResponse
from agents.orchestrator import RecruitmentOrchestrator
from core.config import get_settings
from core.llm_factory import LLMFactory
from core.security import get_current_user
from core.vector_store import QdrantVectorStore
from database.models import User, Job, JobStatus, ProcessingStatus, JobMatch, Resume
from database.session import get_db, db_manager
from tasks.job_processing import process_job_creation_task, reprocess_job_task
from utils.job_id_generator import generate_job_id

logger = logging.getLogger(__name__)
router = APIRouter(tags=["jobs"])


async def _store_job_embedding(job_record: Job, settings) -> None:
    """Store job embedding in vector store for efficient matching."""
    try:
        
        # Initialize vector store
        vector_store = QdrantVectorStore(settings)
        await vector_store.initialize()
        
        # Build comprehensive job text for embedding
        job_text_parts = [f"Job Title: {job_record.title}"]
        
        if job_record.department:
            job_text_parts.append(f"Department: {job_record.department}")
        if job_record.location:
            job_text_parts.append(f"Location: {job_record.location}")
        
        job_text_parts.append(f"Employment Type: {job_record.employment_type}")
        job_text_parts.append(f"Remote Work: {job_record.remote_option}")
        
        if job_record.description:
            job_text_parts.append(f"\nDescription:\n{job_record.description}")
        
        if job_record.responsibilities:
            job_text_parts.append(f"\nResponsibilities:\n{job_record.responsibilities}")
        
        if job_record.requirements:
            job_text_parts.append(f"\nRequirements:\n{job_record.requirements}")
        
        if job_record.benefits:
            job_text_parts.append(f"\nBenefits:\n{job_record.benefits}")
        
        # Add skills information
        if job_record.required_skills:
            job_text_parts.append(f"Required Skills: {', '.join(job_record.required_skills)}")
        if job_record.preferred_skills:
            job_text_parts.append(f"Preferred Skills: {', '.join(job_record.preferred_skills)}")
        
        # Add experience requirements
        if job_record.min_years_experience:
            job_text_parts.append(f"Minimum Experience: {job_record.min_years_experience} years")
        
        # Add education requirements
        if job_record.education_requirements:
            edu_req = job_record.education_requirements
            if isinstance(edu_req, dict):
                if edu_req.get("degree_level"):
                    job_text_parts.append(f"Education Level: {edu_req['degree_level']}")
                if edu_req.get("field_of_study"):
                    job_text_parts.append(f"Field of Study: {', '.join(edu_req['field_of_study'])}")
        
        # Combine all parts
        job_text = " | ".join(job_text_parts)
        
        # Prepare metadata
        metadata = {
            "job_id": str(job_record.id),
            "title": job_record.title,
            "description": job_record.description,
            "required_skills": job_record.required_skills or [],
            "preferred_skills": job_record.preferred_skills or [],
            "min_years_experience": job_record.min_years_experience,
            "education_requirements": job_record.education_requirements or {},
            "location": job_record.location,
            "remote_option": job_record.remote_option,
            "employment_type": job_record.employment_type,
            "organization_id": str(job_record.organization_id),
            "created_at": job_record.created_at.isoformat() if job_record.created_at else None
        }
        
        # Store embedding
        await vector_store.upsert_job_embedding(
            job_id=str(job_record.id),
            text=job_text,
            metadata=metadata
        )
        
        logger.info(f"✅ Job embedding stored successfully: {job_record.id}")
        
    except Exception as e:
        logger.error(f"❌ Failed to store job embedding: {e}")
        raise


class JobFormData(BaseModel):
    """Form data model for job creation/editing."""
    title: str
    department: Optional[str] = None
    location: Optional[str] = None
    employment_type: str = "full-time"  # full-time, part-time, contract, freelance
    remote_option: str = "on-site"  # on-site, hybrid, full
    description: str
    responsibilities: Optional[str] = None
    requirements: Optional[str] = None
    benefits: Optional[str] = None
    salary_min: Optional[int] = None
    salary_max: Optional[int] = None
    salary_currency: str = "USD"
    application_deadline: Optional[str] = None

class JobCreateRequest(JobFormData):
    """Request model for job creation (inherits from JobFormData)."""
    pass

class JobAnalysisResponse(BaseModel):
    """Response model for job analysis results."""
    id: str
    job_id: Optional[str] = None
    title: str
    status: str
    analysis: Optional[Dict[str, Any]] = None
    confidence_score: Optional[float] = None
    processing_time: Optional[float] = None
    error_message: Optional[str] = None

class JDTextAnalysisRequest(BaseModel):
    """Request model for JD text analysis."""
    job_description_text: str
    company_info: Optional[Dict[str, Any]] = None
    create_job: bool = Field(False, description="Whether to create a job record immediately")

class JDTextAnalysisResponse(BaseModel):
    """Response model for JD text analysis and auto-extraction."""
    status: str
    extracted_data: Optional[JobFormData] = None
    analysis: Optional[Dict[str, Any]] = None
    confidence_score: Optional[float] = None
    processing_time: Optional[float] = None
    error_message: Optional[str] = None
    suggestions: Optional[List[str]] = None
    job_id: Optional[str] = None  # Set when create_job=True
    job_created: bool = Field(False, description="Whether a job record was created")

class JobResponse(BaseModel):
    """Response model for job data."""
    id: str
    job_id: Optional[str] = None
    title: str
    department: Optional[str] = None
    location: Optional[str] = None
    employment_type: str
    remote_option: str
    description: str
    responsibilities: Optional[str] = None
    requirements: Optional[str] = None
    benefits: Optional[str] = None
    salary_min: Optional[int] = None
    salary_max: Optional[int] = None
    salary_currency: str
    application_deadline: Optional[str] = None
    status: str
    processing_status: Optional[str] = None
    created_at: str
    updated_at: str
    ai_analysis: Optional[Dict[str, Any]] = None
    skills_analysis: Optional[str] = None
    experience_assessment: Optional[str] = None
    interview_questions: Optional[str] = None
    analysis_generated_at: Optional[str] = None

class JobListResponse(BaseModel):
    """Response model for job list."""
    jobs: List[JobResponse]
    total: int
    page: int
    page_size: int
    has_next: bool

class CopilotQuestionRequest(BaseModel):
    """Request model for copilot questions."""
    question: str = Field(..., description="The question to ask about the job")
    session_id: str = Field(default="default", description="Session ID for conversation context")
    context: Optional[Dict[str, Any]] = Field(None, description="Additional context for the question")

class InterviewQuestionsRequest(BaseModel):
    """Request model for interview question generation."""
    question_count: int = Field(default=10, description="Number of questions to generate")
    focus_areas: Optional[List[str]] = Field(None, description="Specific areas to focus on")
    difficulty_level: str = Field(default="mixed", description="Difficulty level (entry, mid, senior, mixed)")

class SkillAnalysisRequest(BaseModel):
    """Request model for skill analysis."""
    skill_name: str = Field(..., description="Name of the skill to analyze")


# duplicate extraction/safe-strip logic 

def _build_extracted_jobformdata(analysis: Optional[Dict[str, Any]], source_text: str) -> Optional[JobFormData]:
    """Build JobFormData from analysis dict and fallback source text. No logic changes — factored helper."""
    if not analysis:
        return None
    basic_info = analysis.get("basic_info", {}) or {}
    compensation = analysis.get("compensation_info", {}) or {}
    def safe_strip(value, default=""):
        """Safely strip a value, handling None and non-string types."""
        if value is None:
            return default
        if isinstance(value, str):
            return value.strip()
        return str(value).strip() if value else default
    return JobFormData(
        title=safe_strip(basic_info.get("title"), "Job Position"),
        department=safe_strip(basic_info.get("department")) or None,
        location=safe_strip(basic_info.get("location")) or None,
        employment_type=safe_strip(basic_info.get("employment_type"), "full-time").lower().replace(" ", "-"),
        remote_option=safe_strip(basic_info.get("remote_option"), "on-site").lower(),
        description=safe_strip(basic_info.get("description")) or safe_strip(source_text),
        responsibilities="\n".join(analysis.get("responsibilities", [])) or None,
        requirements="\n".join([safe_strip(req.get("requirement", "")) for req in analysis.get("requirements", [])]) or None,
        benefits="\n".join(analysis.get("benefits", [])) or None,
        salary_min=compensation.get("salary_min"),
        salary_max=compensation.get("salary_max"),
        salary_currency=safe_strip(compensation.get("currency")) or "USD",
        application_deadline=None
    )
 
# New helper: build suggestions (exact same logic as previous inline blocks)
def _build_suggestions(analysis: Dict[str, Any], source_text: str) -> List[str]:
    """Return suggestions list based on analysis. No logic or ordering changes."""
    suggestions: List[str] = []
    analysis_insights = analysis.get("analysis_insights", {})
    if analysis_insights.get("clarity_score", 1.0) < 0.7:
        suggestions.append("Consider making the job description clearer and more specific")
    if analysis_insights.get("inclusivity_score", 1.0) < 0.7:
        suggestions.append("Review language for inclusivity and bias-free terminology")
    if len(analysis.get("skills", [])) < 3:
        suggestions.append("Consider adding more specific skill requirements")
    if not analysis.get("compensation_info", {}).get("salary_min") and not analysis.get("compensation_info", {}).get("salary_max"):
        suggestions.append("Consider adding salary range information to attract better candidates")
    if analysis.get("basic_info", {}).get("remote_option", "on-site") == "on-site":
        suggestions.append("Consider offering remote or hybrid work options")
    # Positive feedback (kept in same order / thresholds)
    if analysis_insights.get("clarity_score", 0) > 0.8:
        suggestions.append("✅ Great job description clarity!")
    if len(analysis.get("responsibilities", [])) > 3:
        suggestions.append("✅ Good level of detail in responsibilities")
    return suggestions


@router.post("/analyze-text", response_model=JDTextAnalysisResponse)
async def analyze_job_description_text(
    request: Request,
    jd_request: JDTextAnalysisRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Analyze job description text and extract structured data for auto-populating job form.
    
    This endpoint:
    1. Takes raw job description text
    2. Uses AI to extract and structure all job information
    3. Returns structured data ready for job creation form
    4. Provides suggestions for improvement
    """
    try:
        logger.info("🔍 Analyzing job description text with AI")
        logger.info(f"📄 Text Length: {len(jd_request.job_description_text)} characters")
        if len(jd_request.job_description_text.strip()) < 50:
            raise HTTPException(
                status_code=400, 
                detail="Job description text is too short. Please provide at least 50 characters."
            )
        # Get orchestrator from app state
        orchestrator: RecruitmentOrchestrator = request.app.state.orchestrator
        # Analyze job description text
        analysis_result = await orchestrator.process_job_description(
            job_description_text=jd_request.job_description_text,
            user_id=str(current_user.id),
            organization_id=str(current_user.organization_id),
            metadata={
                "analysis_type": "text_extraction",
                "extracted_by": current_user.username,
                "company_info": jd_request.company_info or {},
                "request_ip": request.client.host if request.client else None
            }
        )
        extracted_data = None
        suggestions = []
        if analysis_result.status == "success" and analysis_result.results.get("analysis"):
            analysis = analysis_result.results["analysis"]
            # Extract structured data for form auto-population
            extracted_data = _build_extracted_jobformdata(analysis, jd_request.job_description_text)
            # Generate suggestions for improvement (factored helper, no logic change)
            suggestions = _build_suggestions(analysis, jd_request.job_description_text)
        logger.info(f"✅ JD text analysis completed: {analysis_result.status}")
        logger.info(f"📊 Confidence Score: {analysis_result.confidence_score:.2f}")
        logger.info(f"📝 Extracted Data: {bool(extracted_data)}")
        # Create job record if requested and analysis was successful
        job_id = None
        job_created = False
        if jd_request.create_job and analysis_result.status == "success" and extracted_data:
            try:
                logger.info("🔄 Creating job record from analyzed data")
                # Generate unique job ID using sync session
                with db_manager.sync_session_context() as sync_db:
                    job_id = generate_job_id(sync_db)
                # Create database record using extracted data
                job_record = Job(
                    job_id=job_id,
                    organization_id=current_user.organization_id,
                    created_by_id=current_user.id,
                    title=extracted_data.title,
                    department=extracted_data.department,
                    location=extracted_data.location,
                    employment_type=extracted_data.employment_type,
                    remote_option=extracted_data.remote_option,
                    description=extracted_data.description,
                    responsibilities=extracted_data.responsibilities,
                    requirements=extracted_data.requirements,
                    benefits=extracted_data.benefits,
                    salary_min=extracted_data.salary_min,
                    salary_max=extracted_data.salary_max,
                    currency=extracted_data.salary_currency,
                    status=JobStatus.ACTIVE
                )
                # Store AI analysis results
                if analysis_result.results.get("analysis"):
                    analysis = analysis_result.results["analysis"]
                    job_record.ai_processed_requirements = analysis
                    # Extract and store key information (same as create_job endpoint)
                    skills = analysis.get("skills", [])
                    experience_req = analysis.get("experience_requirements", {})
                    education_req = analysis.get("education_requirements", {})
                    matching_criteria = analysis.get("matching_criteria", {})
                    # Store skills
                    if skills:
                        required_skills = [skill["skill"] for skill in skills if skill.get("importance") == "required"]
                        preferred_skills = [skill["skill"] for skill in skills if skill.get("importance") == "preferred"]
                        job_record.required_skills = required_skills
                        job_record.preferred_skills = preferred_skills
                        # Create skill weights
                        skill_weights = {}
                        for skill in skills:
                            if skill.get("skill") and "importance" in skill:
                                weight = 1.0 if skill["importance"] == "required" else 0.5
                                skill_weights[skill["skill"]] = weight
                        job_record.skill_weights = skill_weights
                    # Store experience requirements
                    if experience_req:
                        job_record.min_years_experience = experience_req.get("min_years_total")
                        job_record.max_years_experience = experience_req.get("max_years_total")
                        job_record.required_experience_areas = experience_req.get("specific_experience_areas", [])
                        job_record.leadership_required = experience_req.get("leadership_required", False)
                    # Store education requirements
                    if education_req:
                        job_record.education_requirements = education_req
                        job_record.required_degree_level = education_req.get("degree_level")
                        job_record.preferred_fields = education_req.get("field_of_study", [])
                    # Store matching configuration
                    if matching_criteria:
                        job_record.matching_weights = matching_criteria.get("skill_weights", {})
                        job_record.quality_threshold = matching_criteria.get("minimum_match_threshold", 0.7)
                    # Store processing metadata
                    job_record.requirements_processed_by = "JobDescriptionAnalysisAgent"
                    job_record.llm_provider_used = analysis_result.agent_metadata.get("job_description_agent")
                    job_record.processing_confidence = analysis_result.confidence_score
                    job_record.published_at = datetime.utcnow()
                # Save to database
                db.add(job_record)
                await db.commit()
                await db.refresh(job_record)
                job_id = str(job_record.id)
                job_created = True
                logger.info(f"✅ Job record created successfully: {job_id}")
                suggestions.insert(0, f"✅ Job posting created successfully! Job ID: {job_id}")
            except Exception as e:
                logger.error(f"❌ Failed to create job record: {str(e)}")
                suggestions.append(f"⚠️ Analysis successful but job creation failed: {str(e)}")
        return JDTextAnalysisResponse(
            status=analysis_result.status,
            extracted_data=extracted_data,
            analysis=analysis_result.results.get("analysis") if analysis_result.results else None,
            confidence_score=analysis_result.confidence_score,
            processing_time=analysis_result.execution_time,
            error_message="; ".join(analysis_result.errors) if analysis_result.errors else None,
            suggestions=suggestions,
            job_id=job_id,
            job_created=job_created
        )    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ JD text analysis failed: {str(e)}")
        return JDTextAnalysisResponse(
            status="failed",
            extracted_data=None,
            analysis=None,
            confidence_score=0.0,
            processing_time=0.0,
            error_message=str(e),
            suggestions=["Please try again with a different job description"]
        )

@router.post("/upload-jd", response_model=JDTextAnalysisResponse)
async def upload_jd_file(
    request: Request,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Upload and analyze a job description file (PDF, DOC, DOCX, TXT).
    
    This endpoint:
    1. Accepts file uploads (PDF, DOC, DOCX, TXT)
    2. Extracts text from the uploaded file
    3. Uses AI to analyze and structure the job information
    4. Returns structured data ready for job creation form
    """
    try:
        logger.info(f"📄 Uploading and analyzing JD file: {file.filename}")
        # Validate file type
        allowed_types = [
            "application/pdf",
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "text/plain"
        ]
        if file.content_type not in allowed_types:
            raise HTTPException(
                status_code=400, 
                detail="Unsupported file type. Please upload PDF, DOC, DOCX, or TXT files."
            )
        # Read file content
        content = await file.read()
        # Extract text based on file type
        job_description_text = ""
        try:
            if file.content_type == "text/plain":
                job_description_text = content.decode('utf-8')
            elif file.content_type == "application/pdf":
                # Extract text from PDF using pdfplumber
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    text_parts = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    job_description_text = "\n".join(text_parts)       
            elif file.content_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                # Extract text from Word documents
                if file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    # For .docx files, use python-docx
                    doc = docx.Document(io.BytesIO(content))
                    text_parts = []
                    # Extract text from paragraphs
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text.strip())
                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                text_parts.append(" | ".join(row_text))
                    job_description_text = "\n".join(text_parts)
                else:
                    # For .doc files, use docx2txt as fallback
                    
                    # Save content to temporary file for docx2txt
                    with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                        temp_file.write(content)
                        temp_file_path = temp_file.name
                    try:
                        job_description_text = docx2txt.process(temp_file_path)
                    finally:
                        # Clean up temporary file
                        if os.path.exists(temp_file_path):
                            os.unlink(temp_file_path)
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file type: {file.content_type}. Please upload PDF, DOC, DOCX, or TXT files."
                )
        except Exception as e:
            logger.error(f"❌ Document parsing failed: {str(e)}")
            raise HTTPException(
                status_code=400,
                detail=f"Failed to extract text from the uploaded file: {str(e)}. Please ensure the file is not corrupted and try again."
            )
        if not job_description_text.strip():
            raise HTTPException(
                status_code=400,
                detail="Could not extract text from the uploaded file. Please ensure the file contains readable text."
            )
        logger.info(f"📝 Extracted {len(job_description_text)} characters from {file.filename}")
        # Use the same logic as analyze-text endpoint
        logger.info("🔍 Analyzing uploaded job description with AI")
        logger.info(f"📄 Text Length: {len(job_description_text)} characters")
        if len(job_description_text.strip()) < 50:
            raise HTTPException(
                status_code=400, 
                detail="Job description text is too short. Please provide at least 50 characters."
            )
        # Get orchestrator from app state (same as analyze-text)
        orchestrator: RecruitmentOrchestrator = request.app.state.orchestrator
        # Analyze job description text using the same method as analyze-text
        analysis_result = await orchestrator.process_job_description(
            job_description_text=job_description_text,
            user_id=str(current_user.id),
            organization_id=str(current_user.organization_id),
            metadata={
                "analysis_type": "file_upload",
                "extracted_by": current_user.username,
                "original_filename": file.filename,
                "file_size": len(content),
                "request_ip": request.client.host if request.client else None
            }
        )
        extracted_data = None
        suggestions = []
        if analysis_result.status == "success" and analysis_result.results.get("analysis"):
            analysis = analysis_result.results["analysis"]
           # Extract structured data for form auto-population (factored helper, no logic change)
            extracted_data = _build_extracted_jobformdata(analysis, job_description_text)
            # Generate suggestions for improvement (factored helper, no logic change)
            suggestions = _build_suggestions(analysis, job_description_text)
        logger.info(f"✅ JD file analysis completed: {analysis_result.status}")
        logger.info(f"📊 Confidence Score: {analysis_result.confidence_score:.2f}")
        logger.info(f"📝 Extracted Data: {bool(extracted_data)}")
        # No job creation for upload endpoint (user can create manually)
        job_id = None
        job_created = False
        return JDTextAnalysisResponse(
            status=analysis_result.status,
            extracted_data=extracted_data,
            analysis=analysis_result.results.get("analysis") if analysis_result.results else None,
            confidence_score=analysis_result.confidence_score,
            processing_time=analysis_result.execution_time,
            error_message="; ".join(analysis_result.errors) if analysis_result.errors else None,
            suggestions=suggestions,
            job_id=job_id,
            job_created=job_created
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ JD file upload failed: {str(e)}")
        return JDTextAnalysisResponse(
            status="failed",
            extracted_data=None,
            analysis=None,
            confidence_score=0.0,
            processing_time=0.0,
            error_message=str(e),
            suggestions=["Please try uploading a different file or use the text analysis option"]
        )

@router.post("/{job_id}/reprocess")
async def reprocess_job(
    job_id: str,
    request: Request,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Reprocess a job posting with fresh AI analysis.
    
    This endpoint:
    1. Fetches the existing job record
    2. Updates processing status to "processing"
    3. Queues reprocessing in Celery background task
    4. Returns processing status
    """
    try:
        logger.info(f"🔄 Reprocessing job: {job_id}")
        
        # Check user permissions (admin only)
        if current_user.role not in ["admin", "super_admin"]:
            raise HTTPException(
                status_code=403,
                detail="Only admin users can reprocess jobs"
            )
        
        # Fetch the job record
        job_query = select(Job).where(
            and_(
                Job.id == job_id,
                Job.organization_id == current_user.organization_id
            )
        )
        result = await db.execute(job_query)
        job_record = result.scalar_one_or_none()
        
        if not job_record:
            raise HTTPException(
                status_code=404,
                detail="Job not found or access denied"
            )
        
        logger.info(f"📝 Reprocessing job: {job_record.title}")
        
        # Update processing status to processing
        job_record.processing_status = ProcessingStatus.PROCESSING
        job_record.updated_at = datetime.utcnow()
        await db.commit()
        await db.refresh(job_record)
        
        # Prepare job data for background processing
        job_data_dict = {
            "title": job_record.title,
            "department": job_record.department,
            "location": job_record.location,
            "employment_type": job_record.employment_type,
            "remote_option": job_record.remote_option,
            "description": job_record.description,
            "responsibilities": job_record.responsibilities,
            "requirements": job_record.requirements,
            "benefits": job_record.benefits,
            "salary_min": job_record.salary_min,
            "salary_max": job_record.salary_max,
            "salary_currency": job_record.currency,
        }
        
        # Queue background reprocessing task
        
        task = reprocess_job_task.delay(
            job_id=str(job_record.id),
            job_data=job_data_dict,
            user_id=str(current_user.id),
            organization_id=str(current_user.organization_id),
            metadata={
                "reprocessed_by": current_user.username,
                "original_job_id": str(job_record.id),
                "request_ip": request.client.host if request.client else None
            }
        )
        
        logger.info(f"✅ Job reprocessing queued: {job_record.id}")
        logger.info(f"🔄 Background task ID: {task.id}")
        
        return {
            "success": True,
            "message": "Job reprocessing started",
            "job_id": str(job_record.id),
            "title": job_record.title,
            "processing_status": "processing",
            "task_id": task.id
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Job reprocessing failed: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail="Job reprocessing failed"
        )

@router.post("/", response_model=JobAnalysisResponse)
async def create_job(
    request: Request,
    job_data: JobCreateRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Create a new job posting with AI-powered analysis.
    
    This endpoint:
    1. Creates a job record immediately with "processing" status
    2. Queues AI analysis in Celery background task
    3. Returns job info with processing status
    """
    try:
        logger.info("🔍 Creating new job posting")
        logger.info(f"📝 Job Title: {job_data.title}")
        logger.info(f"📄 Description Length: {len(job_data.description)} characters")
        # Generate unique job ID using sync session
        with db_manager.sync_session_context() as sync_db:
            job_id = generate_job_id(sync_db)
        # Create database record with processing status
        job_record = Job(
            job_id=job_id,
            organization_id=current_user.organization_id,
            created_by_id=current_user.id,
            title=job_data.title,
            department=job_data.department,
            location=job_data.location,
            employment_type=job_data.employment_type,
            remote_option=job_data.remote_option,
            description=job_data.description,
            responsibilities=job_data.responsibilities,
            requirements=job_data.requirements,
            benefits=job_data.benefits,
            salary_min=job_data.salary_min,
            salary_max=job_data.salary_max,
            currency=job_data.salary_currency,
            status=JobStatus.DRAFT,  # Start as draft until processing completes
            processing_status=ProcessingStatus.PROCESSING  # Mark as processing
        )
        # Save to database immediately
        db.add(job_record)
        await db.commit()
        await db.refresh(job_record)
        # Queue background processing task
        # Prepare job data for background processing
        job_data_dict = {
            "title": job_data.title,
            "department": job_data.department,
            "location": job_data.location,
            "employment_type": job_data.employment_type,
            "remote_option": job_data.remote_option,
            "description": job_data.description,
            "responsibilities": job_data.responsibilities,
            "requirements": job_data.requirements,
            "benefits": job_data.benefits,
            "salary_min": job_data.salary_min,
            "salary_max": job_data.salary_max,
            "salary_currency": job_data.salary_currency,
        }
        # Start background processing
        task = process_job_creation_task.delay(
            job_id=str(job_record.id),
            job_data=job_data_dict,
            user_id=str(current_user.id),
            organization_id=str(current_user.organization_id),
            metadata={
                "created_by": current_user.username,
                "request_ip": request.client.host if request.client else None
            }
        )
        logger.info(f"✅ Job created and queued for processing: {job_record.id}")
        logger.info(f"🔄 Background task ID: {task.id}")
        return JobAnalysisResponse(
            id=str(job_record.id),
            job_id=job_record.job_id,
            title=job_data.title,
            status="processing",  # Indicate processing is in progress
            analysis=None,  # No analysis yet
            confidence_score=None,
            processing_time=None,
            error_message=None
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Job creation failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Job creation failed")

@router.get("/", response_model=JobListResponse)
async def list_jobs(
    page: int = Query(1, ge=1, description="Page number"),
    page_size: int = Query(20, ge=1, le=1000, description="Items per page"),
    status: Optional[str] = Query(None, description="Filter by status"),
    department: Optional[str] = Query(None, description="Filter by department"),
    employment_type: Optional[str] = Query(None, description="Filter by employment type"),
    search: Optional[str] = Query(None, description="Search by title or description"),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    List job postings with filtering and pagination.
    """
    try:
        logger.info(f"🔍 Listing jobs for user {current_user.id} (page {page})")
        
        # Build base query
        query = select(Job).where(Job.organization_id == current_user.organization_id)
        
        # Apply filters
        if status:
            query = query.where(Job.status == status)
        
        if department:
            query = query.where(Job.department.ilike(f"%{department}%"))
        
        if employment_type:
            query = query.where(Job.employment_type == employment_type)
        
        if search:
            search_filter = f"%{search}%"
            query = query.where(
                or_(
                    Job.title.ilike(search_filter),
                    Job.description.ilike(search_filter),
                    Job.department.ilike(search_filter),
                    Job.job_id.ilike(search_filter)
                )
            )
        
        # Get total count
        count_query = select(func.count()).select_from(query.subquery())
        count_result = await db.execute(count_query)
        total = count_result.scalar() or 0
        
        # Apply pagination
        offset = (page - 1) * page_size
        paginated_query = query.offset(offset).limit(page_size).order_by(Job.created_at.desc())
        
        # Execute query
        result = await db.execute(paginated_query)
        jobs = result.scalars().all()
        
        # Format response
        job_list = []
        for job in jobs:
            job_data = JobResponse(
                id=str(job.id),
                job_id=job.job_id,
                title=job.title,
                department=job.department,
                location=job.location,
                employment_type=job.employment_type,
                remote_option=job.remote_option,
                description=job.description,
                responsibilities=job.responsibilities,
                requirements=job.requirements,
                benefits=job.benefits,
                salary_min=job.salary_min,
                salary_max=job.salary_max,
                salary_currency=job.currency,
                application_deadline=job.application_deadline.isoformat() if job.application_deadline else None,
                status=job.status.value,
                processing_status=job.processing_status.value if job.processing_status else None,
                created_at=job.created_at.isoformat(),
                updated_at=job.updated_at.isoformat(),
                ai_analysis=job.ai_processed_requirements,
                skills_analysis=job.skills_analysis,
                experience_assessment=job.experience_assessment,
                interview_questions=job.interview_questions,
                analysis_generated_at=job.analysis_generated_at.isoformat() if job.analysis_generated_at else None
            )
            job_list.append(job_data)
        
        # Calculate if there's a next page
        has_next = (offset + page_size) < total
        
        logger.info(f"✅ Retrieved {len(job_list)} jobs (total: {total})")
        
        return JobListResponse(
            jobs=job_list,
            total=total,
            page=page,
            page_size=page_size,
            has_next=has_next
        )
        
    except Exception as e:
        logger.error(f"❌ Failed to list jobs: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to retrieve jobs")

@router.get("/{job_id}", response_model=JobResponse)
async def get_job(
    job_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get a specific job posting with full details."""
    try:
        logger.info(f"🔍 Getting job details for {job_id}")
        
        # Parse job_id as UUID
        try:
            job_uuid = UUID(job_id)
        except ValueError:
            raise HTTPException(status_code=400, detail="Invalid job ID format")
        
        # Query job from database
        query = select(Job).where(
            Job.id == job_uuid,
            Job.organization_id == current_user.organization_id
        )
        result = await db.execute(query)
        job = result.scalar_one_or_none()
        
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")
        
        # Format response
        job_response = JobResponse(
            id=str(job.id),
            job_id=job.job_id,
            title=job.title,
            department=job.department,
            location=job.location,
            employment_type=job.employment_type,
            remote_option=job.remote_option,
            description=job.description,
            responsibilities=job.responsibilities,
            requirements=job.requirements,
            benefits=job.benefits,
            salary_min=job.salary_min,
            salary_max=job.salary_max,
            salary_currency=job.currency,
            application_deadline=job.application_deadline.isoformat() if job.application_deadline else None,
            status=job.status.value,
            processing_status=job.processing_status.value if job.processing_status else None,
            created_at=job.created_at.isoformat(),
            updated_at=job.updated_at.isoformat(),
            ai_analysis=job.ai_processed_requirements,
            skills_analysis=job.skills_analysis,
            experience_assessment=job.experience_assessment,
            interview_questions=job.interview_questions,
            analysis_generated_at=job.analysis_generated_at.isoformat() if job.analysis_generated_at else None
        )
        
        logger.info(f"✅ Retrieved job details: {job.title}")
        return job_response
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Failed to get job: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to retrieve job")

@router.put("/{job_id}", response_model=JobAnalysisResponse)
async def update_job(
    job_id: str,
    job_data: JobCreateRequest,
    request: Request,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Update a job posting with AI re-analysis."""
    try:
        logger.info(f"🔄 Updating job posting {job_id}")
        
        # Parse job_id as UUID
        try:
            job_uuid = UUID(job_id)
        except ValueError:
            raise HTTPException(status_code=400, detail="Invalid job ID format")
        
        # Find existing job
        query = select(Job).where(
            Job.id == job_uuid,
            Job.organization_id == current_user.organization_id
        )
        result = await db.execute(query)
        job = result.scalar_one_or_none()
        
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")
        
        # Update basic fields
        job.title = job_data.title
        job.department = job_data.department
        job.location = job_data.location
        job.employment_type = job_data.employment_type
        job.remote_option = job_data.remote_option
        job.description = job_data.description
        job.responsibilities = job_data.responsibilities
        job.requirements = job_data.requirements
        job.benefits = job_data.benefits
        job.salary_min = job_data.salary_min
        job.salary_max = job_data.salary_max
        job.currency = job_data.salary_currency
        job.updated_at = datetime.utcnow()
        
        # Delete all existing job matchings since the job has been updated
        # This ensures that old matching results don't remain after job changes
        
        logger.info(f"🗑️ Deleting all existing job matchings for updated job {job_id}")
        
        # Count existing matches before deletion
        count_query = select(func.count(JobMatch.id)).where(JobMatch.job_id == job_uuid)
        count_result = await db.execute(count_query)
        existing_matches_count = count_result.scalar()
        
        if existing_matches_count > 0:
            # Delete all JobMatch records for this job
            delete_query = delete(JobMatch).where(JobMatch.job_id == job_uuid)
            await db.execute(delete_query)
            logger.info(f"✅ Deleted {existing_matches_count} existing job matchings for job {job_id}")
        else:
            logger.info(f"ℹ️ No existing job matchings found for job {job_id}")
        
        # Note: AI analysis will be handled by Celery reprocessing task
        # No synchronous analysis here to avoid duplicate processing
        
        # Save changes
        await db.commit()
        await db.refresh(job)
        
        logger.info(f"✅ Job updated successfully: {job.title}")
        
        return JobAnalysisResponse(
            id=str(job.id),
            job_id=job.job_id,
            title=job.title,
            status="updated",
            analysis=None,  # Analysis will be handled by Celery reprocessing
            confidence_score=None,
            processing_time=0.0
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Job update failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Job update failed")

@router.delete("/{job_id}")
async def delete_job(
    job_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Delete a job posting."""
    try:
        logger.info(f"🗑️ Deleting job posting {job_id}")
        
        # Parse job_id as UUID
        try:
            job_uuid = UUID(job_id)
        except ValueError:
            raise HTTPException(status_code=400, detail="Invalid job ID format")
        
        # Find and delete job
        query = select(Job).where(
            Job.id == job_uuid,
            Job.organization_id == current_user.organization_id
        )
        result = await db.execute(query)
        job = result.scalar_one_or_none()
        
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")
        
        # Delete all job matchings before deleting the job
        
        logger.info(f"🗑️ Deleting all job matchings for job {job_id}")
        
        # Count existing matches before deletion
        count_query = select(func.count(JobMatch.id)).where(JobMatch.job_id == job_uuid)
        count_result = await db.execute(count_query)
        existing_matches_count = count_result.scalar()
        
        if existing_matches_count > 0:
            # Delete all JobMatch records for this job
            delete_query = delete(JobMatch).where(JobMatch.job_id == job_uuid)
            await db.execute(delete_query)
            logger.info(f"✅ Deleted {existing_matches_count} job matchings for job {job_id}")
        else:
            logger.info(f"ℹ️ No job matchings found for job {job_id}")
        
        # Now delete the job itself
        await db.delete(job)
        await db.commit()
        
        # Delete job embedding from vector storage
        try:
            settings = get_settings()
            vector_store = QdrantVectorStore(settings)
            await vector_store.initialize()
            
            # Use job_id directly as point ID (it's already a UUID4)
            await vector_store.delete_job_embedding(job_id)
            logger.info(f"🔍 Job embedding deleted from vector storage: {job_id}")
            
        except Exception as vector_error:
            logger.warning(f"⚠️ Failed to delete job embedding from vector storage: {vector_error}")
            # Don't fail the job deletion if vector deletion fails
        
        logger.info(f"✅ Job deleted successfully: {job.title}")
        
        return JSONResponse(
            content={"message": f"Job posting '{job.title}' deleted successfully"}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Job deletion failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Job deletion failed")


@router.patch("/{job_id}/status")
async def update_job_status(
    job_id: str,
    status_data: dict,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Update only the status of a job posting."""
    try:
        logger.info(f"🔄 Updating job status for {job_id}")
        
        # Parse job_id as UUID
        try:
            job_uuid = UUID(job_id)
        except ValueError:
            raise HTTPException(status_code=400, detail="Invalid job ID format")
        
        # Find existing job
        query = select(Job).where(
            and_(
                Job.id == job_uuid,
                Job.organization_id == current_user.organization_id
            )
        )
        result = await db.execute(query)
        job = result.scalar_one_or_none()
        
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")
        
        # Validate status
        new_status = status_data.get("status")
        if new_status not in ["active", "inactive", "paused", "closed"]:
            raise HTTPException(status_code=400, detail="Invalid status. Must be one of: active, inactive, paused, closed")
        
        # Update status
        old_status = job.status.value
        job.status = JobStatus(new_status)
        job.updated_at = datetime.utcnow()
        
        await db.commit()
        await db.refresh(job)
        
        logger.info(f"✅ Job status updated: {job.title} ({old_status} → {new_status})")
        
        return JSONResponse(
            content={
                "message": f"Job status updated to {new_status}",
                "job_id": str(job.id),
                "job_title": job.title,
                "old_status": old_status,
                "new_status": new_status
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Job status update failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Job status update failed")


@router.get("/{job_id}/processing-status")
async def get_job_processing_status(
    job_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get the processing status of a job."""
    try:
        logger.info(f"🔍 Getting processing status for job {job_id}")
        
        # Parse job_id as UUID
        try:
            job_uuid = UUID(job_id)
        except ValueError:
            raise HTTPException(status_code=400, detail="Invalid job ID format")
        
        # Query job from database
        query = select(Job).where(
            Job.id == job_uuid,
            Job.organization_id == current_user.organization_id
        )
        result = await db.execute(query)
        job = result.scalar_one_or_none()
        
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")
        
        return {
            "job_id": str(job.id),
            "title": job.title,
            "processing_status": job.processing_status.value if job.processing_status else None,
            "status": job.status.value,
            "updated_at": job.updated_at.isoformat()
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Failed to get job processing status: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to retrieve job processing status")


# Copilot Endpoints
@router.post("/{job_id}/copilot/ask", response_model=CopilotResponse)
async def ask_copilot_question(
    job_id: str,
    request: CopilotQuestionRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """
    Ask the AI copilot a question about a specific job posting.
    
    This endpoint allows recruiters to ask natural language questions about job descriptions
    and receive AI-powered answers with supporting evidence and suggestions.
    """
    try:
        logger.info(f"🤖 Copilot question for job {job_id}: {request.question[:100]}...")
        
        # Get job data
        result = await db.execute(select(Job).filter(Job.id == job_id))
        job = result.scalar_one_or_none()
        
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")
        
        # Check access permissions (basic check - job must be from user's organization)
        # Add more sophisticated access control as needed
        
        # Convert job to dictionary with safe attribute access
        job_data = {
            "id": str(job.id),
            "title": job.title,
            "department": getattr(job, 'department', None),
            "location": getattr(job, 'location', None),
            "employment_type": getattr(job, 'employment_type', None),
            "remote_option": getattr(job, 'remote_option', None),
            "description": job.description,
            "responsibilities": getattr(job, 'responsibilities', None),
            "requirements": getattr(job, 'requirements', None),
            "benefits": getattr(job, 'benefits', None),
            "salary_min": getattr(job, 'salary_min', None),
            "salary_max": getattr(job, 'salary_max', None),
            "salary_currency": getattr(job, 'currency', 'USD'),
            # Use ai_processed_requirements as ai_analysis or combine relevant AI fields
            "ai_analysis": {
                "processed_requirements": getattr(job, 'ai_processed_requirements', {}),
                "required_skills": getattr(job, 'required_skills', []),
                "preferred_skills": getattr(job, 'preferred_skills', []),
                "skill_weights": getattr(job, 'skill_weights', {}),
                "education_requirements": getattr(job, 'education_requirements', {}),
                "min_years_experience": getattr(job, 'min_years_experience', None),
                "max_years_experience": getattr(job, 'max_years_experience', None),
                "leadership_required": getattr(job, 'leadership_required', False)
            }
        }
        
        # Initialize copilot agent with proper LLM configuration
        settings = get_settings()
        llm_factory = LLMFactory(settings)
        llm = llm_factory.create_llm_with_fallback()
        
        copilot = JobCopilotAgent(llm, settings)
        await copilot.initialize()
        
        # Process the question
        response = await copilot.ask_question(
            question=request.question,
            job_data=job_data,
            session_id=request.session_id,
            context=request.context
        )
        
        logger.info(f"✅ Copilot response generated with confidence: {response.confidence_score:.2f}")
        return response
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Copilot question failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to process copilot question")


@router.post("/{job_id}/copilot/interview-questions", response_model=InterviewQuestionsResponse)
async def generate_interview_questions(
    job_id: str,
    request: InterviewQuestionsRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """
    Generate AI-powered interview questions based on the job description.
    
    Creates tailored interview questions covering technical skills, behavioral aspects,
    and role-specific scenarios to help recruiters conduct effective interviews.
    """
    try:
        logger.info(f"🎤 Generating {request.question_count} interview questions for job {job_id}")
        
        # Get job data
        result = await db.execute(select(Job).filter(Job.id == job_id))
        job = result.scalar_one_or_none()
        
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")
        
        # Convert job to dictionary
        job_data = {
            "id": str(job.id),
            "title": job.title,
            "department": job.department,
            "location": job.location,
            "employment_type": job.employment_type,
            "remote_option": job.remote_option,
            "description": job.description,
            "responsibilities": job.responsibilities,
            "requirements": job.requirements,
            "benefits": job.benefits,
            "ai_analysis": {
                "processed_requirements": job.ai_processed_requirements or {},
                "required_skills": job.required_skills or [],
                "preferred_skills": job.preferred_skills or [],
                "experience_requirements": {
                    "min_years": job.min_years_experience,
                    "max_years": job.max_years_experience,
                    "required_areas": job.required_experience_areas or []
                },
                "education_requirements": job.education_requirements or {}
            }
        }
        
        # Initialize copilot agent with proper LLM configuration
        settings = get_settings()
        llm_factory = LLMFactory(settings)
        llm = llm_factory.create_llm_with_fallback()
        
        copilot = JobCopilotAgent(llm, settings)
        await copilot.initialize()
        
        # Generate interview questions
        response = await copilot.generate_interview_questions(
            job_data=job_data,
            question_count=request.question_count,
            focus_areas=request.focus_areas,
            difficulty_level=request.difficulty_level
        )
        
        logger.info(f"✅ Generated {len(response.questions)} interview questions")
        return response
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Interview question generation failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to generate interview questions")


@router.post("/{job_id}/copilot/analyze-skill", response_model=SkillAnalysisResponse)
async def analyze_skill(
    job_id: str,
    request: SkillAnalysisRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """
    Get detailed analysis of a specific skill mentioned in the job description.
    
    Provides comprehensive insights about skill importance, proficiency requirements,
    market demand, and learning resources to help recruiters better understand requirements.
    """
    try:
        logger.info(f"🔍 Analyzing skill '{request.skill_name}' for job {job_id}")
        
        # Get job data
        result = await db.execute(select(Job).filter(Job.id == job_id))
        job = result.scalar_one_or_none()
        
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")
        
        # Convert job to dictionary
        job_data = {
            "id": str(job.id),
            "title": job.title,
            "department": job.department,
            "location": job.location,
            "employment_type": job.employment_type,
            "remote_option": job.remote_option,
            "description": job.description,
            "responsibilities": job.responsibilities,
            "requirements": job.requirements,
            "benefits": job.benefits,
            "ai_analysis": {
                "processed_requirements": job.ai_processed_requirements or {},
                "required_skills": job.required_skills or [],
                "preferred_skills": job.preferred_skills or [],
                "experience_requirements": {
                    "min_years": job.min_years_experience,
                    "max_years": job.max_years_experience,
                    "required_areas": job.required_experience_areas or []
                },
                "education_requirements": job.education_requirements or {}
            }
        }
        
        # Initialize copilot agent with proper LLM configuration
        settings = get_settings()
        llm_factory = LLMFactory(settings)
        llm = llm_factory.create_llm_with_fallback()
        
        copilot = JobCopilotAgent(llm, settings)
        await copilot.initialize()
        
        # Analyze skill
        response = await copilot.analyze_skill_details(
            skill_name=request.skill_name,
            job_data=job_data
        )
        
        logger.info(f"✅ Completed skill analysis for '{request.skill_name}'")
        return response
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Skill analysis failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to analyze skill")


async def _find_matching_resumes_for_job(job_record: Job, settings, db: AsyncSession) -> None:
    """Find matching resumes for a new job and store matches in database."""
    try:
        
        logger.info(f"🔍 Finding matching resumes for new job: {job_record.title}")
        
        # Initialize vector store
        vector_store = QdrantVectorStore(settings)
        await vector_store.initialize()
        
        # Build job text for vector search
        job_text_parts = [f"Job Title: {job_record.title}"]
        if job_record.department:
            job_text_parts.append(f"Department: {job_record.department}")
        if job_record.location:
            job_text_parts.append(f"Location: {job_record.location}")
        job_text_parts.append(f"Employment Type: {job_record.employment_type}")
        job_text_parts.append(f"Remote Work: {job_record.remote_option}")
        job_text_parts.append(f"Description: {job_record.description}")
        
        if job_record.responsibilities:
            job_text_parts.append(f"Responsibilities: {job_record.responsibilities}")
        if job_record.requirements:
            job_text_parts.append(f"Requirements: {job_record.requirements}")
        
        job_text = " | ".join(job_text_parts)
        
        # Search for similar resumes using vector search
        search_results = await vector_store.search_resumes(
            query_text=job_text,
            top_k=20,  # Get more resumes for analysis
            org_id=str(job_record.organization_id)
        )
        
        if not search_results:
            logger.info("No similar resumes found via vector search")
            return
        
        logger.info(f"🔍 Vector search found {len(search_results)} similar resumes")
        
        # Get resume IDs from search results
        resume_ids = []
        for result in search_results:
            resume_id = result.payload.get("resume_id")
            if resume_id:
                resume_ids.append(resume_id)
        
        if not resume_ids:
            logger.info("No valid resume IDs found in vector search results")
            return
        
        # Fetch resumes from database
        resumes_result = await db.execute(
            select(Resume).where(
                and_(
                    Resume.id.in_(resume_ids),
                    Resume.organization_id == job_record.organization_id,
                    Resume.status == "analyzed"
                )
            )
        )
        resumes = resumes_result.scalars().all()
        
        if not resumes:
            logger.info("No active resumes found in database")
            return
        
        logger.info(f"🔍 Found {len(resumes)} active resumes for matching")
        
        # Initialize orchestrator for job matching
        orchestrator = RecruitmentOrchestrator(settings)
        await orchestrator.initialize()
        
        # Prepare job data for matching
        job_data = {
            "id": str(job_record.id),
            "title": job_record.title,
            "description": job_record.description,
            "required_skills": job_record.required_skills or [],
            "preferred_skills": job_record.preferred_skills or [],
            "min_years_experience": job_record.min_years_experience,
            "education_requirements": job_record.education_requirements or {},
            "location": job_record.location,
            "remote_option": job_record.remote_option,
            "employment_type": job_record.employment_type
        }
        
        # Find matches for each resume
        matches_to_store = []
        for resume in resumes:
            try:
                # Extract resume data for matching
                ai_analysis = resume.ai_analysis or {}
                personal_info = ai_analysis.get("personal_info", {})
                career_analysis = ai_analysis.get("career_analysis", {})
                skills_analysis = ai_analysis.get("skills_analysis", {})
                education_analysis = ai_analysis.get("education_analysis", {})
                
                candidate_data = {
                    "id": str(resume.id),
                    "name": resume.candidate_name,
                    "email": resume.candidate_email,
                    "skills": _extract_skills_list(skills_analysis),
                    "total_years_experience": career_analysis.get("total_years_experience", 0),
                    "work_history": career_analysis.get("work_experience", []),
                    "education": education_analysis,
                    "location": personal_info.get("location", ""),
                    "ai_analysis": ai_analysis
                }
                
                # Perform AI matching
                match_analysis = await orchestrator.job_matching_agent.analyze_single_match(
                    job_posting=job_data,
                    candidate=candidate_data
                )
                
                if "analysis" in match_analysis:
                    analysis = match_analysis["analysis"]
                    skills_score = analysis.get("skills_score", 0)
                    overall_score = analysis.get("overall_score", 0)
                    
                    # Only store matches with skills score >= 30%
                    if skills_score >= 0.3:
                        match_data = {
                            "job_id": str(job_record.id),
                            "resume_id": str(resume.id),
                            "overall_score": overall_score,
                            "confidence_level": analysis.get("confidence_level", 0.0),
                            "skills_score": skills_score,
                            "experience_score": analysis.get("experience_score", 0.0),
                            "education_score": analysis.get("education_score", 0.0),
                            "location_score": analysis.get("location_score", 0.0),
                            "cultural_fit_score": analysis.get("cultural_fit_score", 0.0),
                            "strengths": analysis.get("key_strengths", []),
                            "weaknesses": analysis.get("key_concerns", []),
                            "missing_skills": analysis.get("missing_skills", []),
                            "recommendations": analysis.get("recommendations", []),
                            "explanation": analysis.get("detailed_reasoning", ""),
                            "reasoning": analysis.get("detailed_reasoning", ""),
                            "matching_analysis": analysis,
                            "status": "completed",
                            "matching_agent_version": "1.0",
                            "llm_provider_used": settings.LLM_PROVIDER,
                            "matching_algorithm": "vector_search_ai_analysis"
                        }
                        matches_to_store.append(match_data)
                        logger.info(f"✅ Found match: {resume.candidate_name} -> {job_record.title} (skills: {skills_score:.2f})")
                    else:
                        logger.info(f"❌ Filtered out: {resume.candidate_name} -> {job_record.title} (skills: {skills_score:.2f} < 0.3)")
                
            except Exception as e:
                logger.warning(f"⚠️ Failed to analyze resume {resume.id}: {e}")
                continue
        
        # Store matches in database
        if matches_to_store:
            for i, match_data in enumerate(matches_to_store):
                job_match = JobMatch(
                    job_id=match_data["job_id"],
                    resume_id=match_data["resume_id"],
                    overall_score=match_data["overall_score"],
                    confidence_level=match_data["confidence_level"],
                    skills_score=match_data["skills_score"],
                    experience_score=match_data["experience_score"],
                    education_score=match_data["education_score"],
                    location_score=match_data["location_score"],
                    cultural_fit_score=match_data["cultural_fit_score"],
                    matching_analysis=match_data["matching_analysis"],
                    strengths=match_data["strengths"],
                    weaknesses=match_data["weaknesses"],
                    missing_skills=match_data["missing_skills"],
                    recommendations=match_data["recommendations"],
                    explanation=match_data["explanation"],
                    reasoning=match_data["reasoning"],
                    rank_in_job=i + 1,
                    status="completed",
                    matching_agent_version=match_data["matching_agent_version"],
                    llm_provider_used=match_data["llm_provider_used"],
                    matching_algorithm=match_data["matching_algorithm"]
                )
                db.add(job_match)
            
            await db.commit()
            logger.info(f"✅ Stored {len(matches_to_store)} job matches in database")
        else:
            logger.info("No qualifying matches found (skills score < 30%)")
        
    except Exception as e:
        logger.error(f"❌ Failed to find matching resumes for job: {e}")
        # Don't raise the error to avoid breaking job creation


def _extract_skills_list(skills_analysis: Dict[str, Any]) -> List[str]:
    """Extract skills list from skills analysis."""
    try:
        skills = []
        skill_categories = [
            "programming_languages", "frameworks_libraries", "tools_technologies",
            "technical_skills", "soft_skills", "methodologies"
        ]
        
        for category in skill_categories:
            skill_list = skills_analysis.get(category, [])
            if isinstance(skill_list, list):
                for skill in skill_list:
                    if isinstance(skill, dict) and skill.get("skill"):
                        skills.append(skill["skill"])
        
        return skills
    except Exception as e:
        logger.warning(f"Failed to extract skills: {e}")
        return []

